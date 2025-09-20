#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
import sys

def examine_docx(filename):
    print(f"\n=== Examining {filename} ===")
    try:
        doc = Document(filename)
        
        print(f"Number of paragraphs: {len(doc.paragraphs)}")
        print("\nContent:")
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                print(f"{i+1}: {paragraph.text}")
        
        print(f"\nNumber of tables: {len(doc.tables)}")
        for i, table in enumerate(doc.tables):
            print(f"\nTable {i+1}:")
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                print(f"  Row {row_idx+1}: {' | '.join(row_data)}")
                
    except Exception as e:
        print(f"Error reading {filename}: {e}")

if __name__ == "__main__":
    examine_docx("PREVENTIVO BARBERI LACMUS(2).docx")
    examine_docx("template_preventivo.dotx")
