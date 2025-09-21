#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import streamlit as st
import pandas as pd
from datetime import datetime, date
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import os
import json

# Configurazione della pagina per mobile
st.set_page_config(
    page_title="Generatore Preventivi Catering",
    page_icon="🍽️",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# CSS personalizzato per mobile
st.markdown("""
<style>
    .main > div {
        padding-top: 2rem;
    }
    .stButton > button {
        width: 100%;
        margin-top: 10px;
    }
    .stSelectbox > div > div {
        font-size: 16px;
    }
    .stTextInput > div > div > input {
        font-size: 16px;
    }
    .stTextArea > div > div > textarea {
        font-size: 16px;
    }
    @media (max-width: 768px) {
        .main .block-container {
            padding-left: 1rem;
            padding-right: 1rem;
        }
    }
</style>
""", unsafe_allow_html=True)

def initialize_session_state():
    """Inizializza lo stato della sessione"""
    if 'menu_items' not in st.session_state:
        st.session_state.menu_items = []
    if 'quote_data' not in st.session_state:
        st.session_state.quote_data = {}

def load_menu_items():
    """Carica gli elementi del menu dal file JSON"""
    menu_file = "/home/grigolet/cernbox/personal/code/homelette-preventivi/menu_items.json"
    
    # Menu predefiniti di base
    default_items = {
        "COCKTAIL DI BENVENUTO": {
            "categoria": "Aperitivo",
            "descrizione": "Cocktail di benvenuto per gli ospiti"
        },
        "TAGLIATA DI FRUTTA FRESCA": {
            "categoria": "Frutta",
            "descrizione": "Selezione di frutta fresca di stagione"
        },
        "MACEDONIA": {
            "categoria": "Frutta",
            "descrizione": "Macedonia di frutta fresca"
        },
        "PROSCIUTTO DI PARMA E MELONE SU PANE AI CEREALI": {
            "categoria": "Antipasti",
            "descrizione": "Prosciutto di Parma e melone serviti su pane ai cereali"
        },
        "PINZIMONIO DI VERDURE CON SALSE MISTE": {
            "categoria": "Antipasti",
            "descrizione": "Verdure fresche con selezione di salse"
        },
        "COUPELLE DI BRISE' CON MOUSSE DI BOLOGNA E PISTACCHI": {
            "categoria": "Pasticceria Salata",
            "descrizione": "Pasta brise' (farina di frumento, glutine, acqua, burro), mortadella, formaggio fresco, pistacchi"
        },
        "CESTINI DI SFOGLIA CON PATE' DI FAGIOLINI E TONNO AL LIMONE": {
            "categoria": "Pasticceria Salata",
            "descrizione": "Farina 00, burro, panna, gelatina, fagiolini, tonno, limone, sale, pepe"
        },
        "FIORI DI PANE AI CEREALI CON MOUSSE AL PROSCIUTTO E PEPE ROSA": {
            "categoria": "Pasticceria Salata",
            "descrizione": "Farina di cereali, acqua, lievito, sale, prosciutto cotto, formaggio fresco, sale, pepe"
        },
        "FIORI DI PANE ALLA CURCUMA CON PATE' GUSTOSO": {
            "categoria": "Pasticceria Salata",
            "descrizione": "Farina di frumento, acqua, sale, curcuma, fagioli, speck, aglio, rosmarino"
        },
        "FIORI DI PANE AL BASILICO CON PESTO AI POMODORI SECCHI E MANDORLE": {
            "categoria": "Pasticceria Salata",
            "descrizione": "Farina di frumento, basilico, acqua, sale, pomodori secchi, parmigiano, mandorle, basilico, capperi, olio d'oliva"
        },
        "FIORI DI PANE ALLA BARBABIETOLA CON MOUSSE AI 4 FORMAGGI ED ERBE AROMATICHE": {
            "categoria": "Pasticceria Salata",
            "descrizione": "Farina di frumento, basilico, barbabietola, acqua, sale, zola, taleggio, parmigiano, formaggio fresco, erbe, pepe"
        },
        "GIRO FOCACCIA MORBIDA ALLE PATATE": {
            "categoria": "Pane",
            "descrizione": "Farina 0, fiocchi di patate, lievito birra, miele, olio evo, semola, pomodorini, olive"
        },
        "SGABEI ALLE OLIVE": {
            "categoria": "Pane",
            "descrizione": "Farina 0, fiocchi di patate, lievito birra, olive, sale"
        },
        "POMODORINI CROCCANTI": {
            "categoria": "Contorni",
            "descrizione": "Pomodorini, caramello, semi di sesamo"
        },
        "VITELLO TONNATO E CAPPERI": {
            "categoria": "Secondi",
            "descrizione": "Vitello tonnato con capperi"
        },
        "CAPONATA": {
            "categoria": "Contorni",
            "descrizione": "Caponata siciliana"
        },
        "INVOLTINI DI ZUCCHINE": {
            "categoria": "Contorni",
            "descrizione": "Involtini di zucchine ripieni"
        },
        "KEDGEREE AL SALMONE": {
            "categoria": "Primi",
            "descrizione": "Kedgeree al salmone"
        },
        "INSALATA DI PASTA MEDITERRANEA": {
            "categoria": "Primi",
            "descrizione": "Insalata di pasta con ingredienti mediterranei"
        },
        "FANTASIA DI TIRAMISU'": {
            "categoria": "Dolci",
            "descrizione": "Tiramisu' della casa"
        },
        "SPUMONE ALLA FRUTTA": {
            "categoria": "Dolci",
            "descrizione": "Spumone con frutta fresca"
        },
        "CREMA CAFFE'": {
            "categoria": "Dolci",
            "descrizione": "Crema al caffè"
        }
    }
    
    try:
        if os.path.exists(menu_file):
            with open(menu_file, 'r', encoding='utf-8') as f:
                loaded_items = json.load(f)
                # Merge con gli elementi predefiniti, dando priorità a quelli salvati
                default_items.update(loaded_items)
        return default_items
    except Exception as e:
        st.error(f"Errore nel caricamento del menu: {e}")
        return default_items

def save_menu_items(items):
    """Salva gli elementi del menu nel file JSON"""
    menu_file = "/home/grigolet/cernbox/personal/code/homelette-preventivi/menu_items.json"
    try:
        with open(menu_file, 'w', encoding='utf-8') as f:
            json.dump(items, f, ensure_ascii=False, indent=2)
    except Exception as e:
        st.error(f"Errore nel salvataggio del menu: {e}")

def add_new_menu_item(nome, categoria, descrizione):
    """Aggiunge un nuovo elemento al menu persistente"""
    menu_items = load_menu_items()
    menu_items[nome.upper()] = {
        "categoria": categoria,
        "descrizione": descrizione
    }
    save_menu_items(menu_items)
    return menu_items

def create_word_document(quote_data):
    """Crea un documento Word con il preventivo"""
    doc = Document()
    
    # Imposta margini
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Logo e intestazione aziendale
    try:
        logo_path = "/home/grigolet/cernbox/personal/code/homelette-preventivi/homelette_logo.jpg"
        if os.path.exists(logo_path):
            logo_para = doc.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            logo_run = logo_para.runs[0] if logo_para.runs else logo_para.add_run()
            logo_run.add_picture(logo_path, width=Inches(2))
        else:
            # Se il logo non esiste, aggiungi solo il testo del nome
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run("HOMELETTE")
            title_run.font.size = Pt(28)
            title_run.font.bold = True
            title_run.font.name = "Arial Black"  # Bold sans-serif heading font
    except Exception:
        # Fallback se ci sono problemi con l'immagine
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run("HOMELETTE")
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.name = "Arial Black"  # Bold sans-serif heading font
    
    # Subtitle in cobalt blue
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run("Primum manducare, deinde filosofare")
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.bold = True
    subtitle_run.font.name = "Calibri"  # Clean sans-serif for subtitle
    # Imposta colore blu cobalto (RGB: 0, 71, 171)
    subtitle_run.font.color.rgb = RGBColor(0, 71, 171)
    
    # Spazio dopo il subtitle
    doc.add_paragraph()
    
    # Linea decorativa orizzontale
    divider_para = doc.add_paragraph()
    divider_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    divider_run = divider_para.add_run("━" * 60)
    divider_run.font.color.rgb = RGBColor(0, 71, 171)
    
    # Intestazione del preventivo
    header = doc.add_paragraph()
    ref_run = header.add_run("Referenza: ")
    ref_run.bold = True
    ref_run.font.size = Pt(12)
    ref_run.font.name = "Segoe UI"  # Sans-serif for labels
    
    ref_value = header.add_run(quote_data.get('riferimento', ''))
    ref_value.font.name = "Georgia"  # Serif for variable text
    ref_value.font.size = Pt(11)
    
    # Altri campi con font appropriati
    luogo_para = doc.add_paragraph()
    luogo_label = luogo_para.add_run("Luogo: ")
    luogo_label.font.name = "Segoe UI"
    luogo_label.font.size = Pt(12)
    luogo_label.font.bold = True
    luogo_value = luogo_para.add_run(quote_data.get('luogo', ''))
    luogo_value.font.name = "Georgia"
    luogo_value.font.size = Pt(11)
    
    ora_para = doc.add_paragraph()
    ora_label = ora_para.add_run("Ora e Luogo: ")
    ora_label.font.name = "Segoe UI"
    ora_label.font.size = Pt(12)
    ora_label.font.bold = True
    ora_value = ora_para.add_run(quote_data.get('data_ora', ''))
    ora_value.font.name = "Georgia"
    ora_value.font.size = Pt(11)
    
    tipo_para = doc.add_paragraph()
    tipo_label = tipo_para.add_run("Tipologia servizio: ")
    tipo_label.font.name = "Segoe UI"
    tipo_label.font.size = Pt(12)
    tipo_label.font.bold = True
    tipo_value = tipo_para.add_run(quote_data.get('tipologia_servizio', ''))
    tipo_value.font.name = "Georgia"
    tipo_value.font.size = Pt(11)
    
    # Linea separatrice
    divider_para = doc.add_paragraph()
    divider_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    divider_run = divider_para.add_run("━" * 60)
    divider_run.font.color.rgb = RGBColor(0, 71, 171)
    
    doc.add_paragraph()
    
    # Saluto
    saluto = doc.add_paragraph()
    saluto_run = saluto.add_run(f"Gentile {quote_data.get('destinatario', '')},")
    saluto_run.font.name = "Segoe UI"
    saluto_run.font.size = Pt(11)
    
    intro = doc.add_paragraph()
    intro_run = intro.add_run("come da accordi Vi proponiamo la nostra migliore offerta per l'evento di cui in oggetto per numero ")
    intro_run.font.name = "Segoe UI"
    intro_run.font.size = Pt(11)
    numero_run = intro.add_run(str(quote_data.get('numero_persone', '')))
    numero_run.font.name = "Georgia"
    numero_run.font.size = Pt(11)
    intro_run2 = intro.add_run(" persone.")
    intro_run2.font.name = "Segoe UI"
    intro_run2.font.size = Pt(11)
    
    doc.add_paragraph()
    
    # Dettagli servizio con formattazione migliorata
    def add_detail_paragraph(label, value):
        para = doc.add_paragraph()
        label_run = para.add_run(f"{label}: ")
        label_run.font.name = "Segoe UI"
        label_run.font.size = Pt(11)
        label_run.font.bold = True
        value_run = para.add_run(value)
        value_run.font.name = "Georgia"
        value_run.font.size = Pt(11)
    
    add_detail_paragraph("Tipologia del servizio", quote_data.get('tipologia_buffet', 'Buffet e servizio'))
    add_detail_paragraph("Bicchieri", quote_data.get('bicchieri', 'Calici da vino in materiali ecocompatibili usa e getta'))
    add_detail_paragraph("Posate", quote_data.get('posate', 'Acciaio leggero'))
    add_detail_paragraph("Stoviglie", quote_data.get('stoviglie', 'Ceramica e materiali ecocompatibili usa e getta'))
    add_detail_paragraph("Acqua", quote_data.get('acqua', 'Esclusa / da definire'))
    add_detail_paragraph("Vini", quote_data.get('vini', 'Esclusi / da definire'))
    add_detail_paragraph("Servizio", quote_data.get('servizio', 'Incluso'))
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Valorizzazione dell'offerta
    valorizzazione = doc.add_paragraph()
    valorizzazione_run = valorizzazione.add_run("Valorizzazione dell'offerta")
    valorizzazione_run.font.size = Pt(16)
    valorizzazione_run.font.bold = True
    valorizzazione_run.font.name = "Calibri"  # Sans-serif for headings
    
    prezzo_persona = quote_data.get('prezzo_persona', 0)
    numero_persone = quote_data.get('numero_persone', 0)
    
    # Paragrafo con testo statico e variabile
    costo_para = doc.add_paragraph()
    costo_run1 = costo_para.add_run("Il costo del servizio è di ")
    costo_run1.font.name = "Segoe UI"
    costo_run1.font.size = Pt(11)
    
    prezzo_run = costo_para.add_run(f"{prezzo_persona}")
    prezzo_run.font.name = "Georgia"
    prezzo_run.font.size = Pt(11)
    prezzo_run.font.bold = True
    
    costo_run2 = costo_para.add_run(" Euro/persona per un minimo di ")
    costo_run2.font.name = "Segoe UI"
    costo_run2.font.size = Pt(11)
    
    persone_run = costo_para.add_run(f"{numero_persone}")
    persone_run.font.name = "Georgia"
    persone_run.font.size = Pt(11)
    persone_run.font.bold = True
    
    costo_run3 = costo_para.add_run(" persone.")
    costo_run3.font.name = "Segoe UI"
    costo_run3.font.size = Pt(11)
    
    # Paragrafi informativi
    info_para1 = doc.add_paragraph()
    info_run1 = info_para1.add_run("A fronte del saldo della prestazione mediante bonifico bancario verrà emessa su richiesta fattura.")
    info_run1.font.name = "Segoe UI"
    info_run1.font.size = Pt(11)
    
    info_para2 = doc.add_paragraph()
    info_run2 = info_para2.add_run("Resta a nostro carico la preparazione dei tavoli, il riordino del locale e lo smaltimento rifiuti (no vetro)")
    info_run2.font.name = "Segoe UI"
    info_run2.font.size = Pt(11)
    
    doc.add_paragraph()
    
    # Calcoli
    totale_base = prezzo_persona * numero_persone
    costo_cameriere = quote_data.get('costo_cameriere', 0)
    totale_finale = totale_base + costo_cameriere
    
    totale_para = doc.add_paragraph()
    totale_run = totale_para.add_run(f"      Totale {numero_persone} convenuti:        € {totale_base}")
    totale_run.font.size = Pt(12)
    totale_run.font.name = "Segoe UI"
    
    if costo_cameriere > 0:
        doc.add_paragraph()
        cameriere_para = doc.add_paragraph()
        cameriere_run = cameriere_para.add_run(f"      Cameriere                       €   {costo_cameriere}")
        cameriere_run.font.size = Pt(12)
        cameriere_run.font.name = "Segoe UI"
    
    doc.add_paragraph()
    totale_finale_para = doc.add_paragraph()
    totale_finale_run = totale_finale_para.add_run("TOTALE (*)                          ")
    totale_finale_run.bold = True
    totale_finale_run.font.size = Pt(16)
    totale_finale_run.font.name = "Calibri"
    totale_prezzo_run = totale_finale_para.add_run(f"€ {totale_finale}")
    totale_prezzo_run.bold = True
    totale_prezzo_run.font.size = Pt(16)
    totale_prezzo_run.font.name = "Georgia"
    
    # Nota iva
    iva_para = doc.add_paragraph()
    iva_run = iva_para.add_run("(*) Escluso iva 22%")
    iva_run.font.name = "Segoe UI"
    iva_run.font.size = Pt(10)
    
    # Add page break before MENU section
    doc.add_page_break()
    
    # Subtitle in cobalt blue
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run("MENU")
    subtitle_run.font.size = Pt(18)
    subtitle_run.font.bold = True
    subtitle_run.font.name = "Calibri"  # Clean sans-serif for subtitle
    # Imposta colore blu cobalto (RGB: 0, 71, 171)
    subtitle_run.font.color.rgb = RGBColor(0, 71, 171)
    
    # Linea decorativa orizzontale
    divider_para = doc.add_paragraph()
    divider_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    divider_run = divider_para.add_run("━" * 60)
    divider_run.font.color.rgb = RGBColor(0, 71, 171)
    
    
    
    # Menu - Grouped by category
    if quote_data.get('menu_items'):
        # Group items by category
        menu_by_category = {}
        for item in quote_data.get('menu_items', []):
            categoria = item.get('categoria', 'Altro')
            if categoria not in menu_by_category:
                menu_by_category[categoria] = []
            menu_by_category[categoria].append(item)
        
        # Define category order for better presentation
        category_order = [
            'Aperitivo', 'Antipasti', 'Primi', 'Secondi', 
            'Contorni', 'Pasticceria Salata', 'Pane', 
            'Dolci', 'Frutta', 'Bevande', 'Altro'
        ]
        
        # Display menu items grouped by category
        for categoria in category_order:
            if categoria in menu_by_category and menu_by_category[categoria]:
                # Category header
                category_para = doc.add_paragraph()
                category_run = category_para.add_run(categoria.upper())
                category_run.bold = True
                category_run.font.size = Pt(14)
                category_run.font.name = "Calibri"
                category_run.font.color.rgb = RGBColor(0, 71, 171)  # Cobalt blue
                
                # Items in this category
                for item in menu_by_category[categoria]:
                    item_para = doc.add_paragraph()
                    item_run = item_para.add_run(item['nome'].upper())
                    item_run.bold = True
                    item_run.font.size = Pt(12)
                    item_run.font.name = "Calibri"  # Sans-serif for menu headings
                    if item.get('descrizione'):
                        desc_para = doc.add_paragraph(item['descrizione'])
                        desc_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        for run in desc_para.runs:
                            run.font.name = "Segoe UI"
                            run.font.size = Pt(10)
                
                # Add space between categories
                doc.add_paragraph()
    
    # Note aggiuntive
    if quote_data.get('note'):
        doc.add_paragraph()
        doc.add_paragraph()
        note_para = doc.add_paragraph()
        note_run = note_para.add_run("NOTE AGGIUNTIVE:")
        note_run.bold = True
        note_run.font.size = Pt(14)
        note_run.font.name = "Calibri"  # Sans-serif for heading
        
        note_content_para = doc.add_paragraph(quote_data['note'])
        for run in note_content_para.runs:
            run.font.name = "Georgia"  # Serif for variable content
            run.font.size = Pt(11)
    
    return doc

def main():
    initialize_session_state()
    
    st.title("🍽️ Generatore Preventivi Catering")
    st.markdown("*Crea preventivi professionali per eventi di catering*")
    
    # Sidebar per navigazione mobile
    with st.sidebar:
        st.markdown("### Menu di Navigazione")
        sezione = st.radio(
            "Seleziona sezione:",
            ["📋 Dati Evento", "🍽️ Menu", "💰 Prezzi", "📄 Anteprima"]
        )
    
    if sezione == "📋 Dati Evento":
        st.header("📋 Informazioni Evento")
        
        col1, col2 = st.columns([1, 1])
        
        with col1:
            riferimento = st.text_input("Riferimento Cliente", 
                                      value=st.session_state.quote_data.get('riferimento', ''),
                                      placeholder="es. Barberi Mauro")
            
            destinatario = st.text_input("Destinatario", 
                                       value=st.session_state.quote_data.get('destinatario', ''),
                                       placeholder="es. Mauro")
            
            luogo = st.text_input("Luogo Evento", 
                                value=st.session_state.quote_data.get('luogo', ''),
                                placeholder="es. Villa Rondinella")
        
        with col2:
            data_evento = st.date_input("Data Evento", 
                                      value=st.session_state.quote_data.get('data_evento', date.today()))
            
            ora_evento = st.time_input("Ora Evento", 
                                     value=st.session_state.quote_data.get('ora_evento', datetime.now().time()))
            
            numero_persone = st.number_input("Numero Persone", 
                                           min_value=1, 
                                           value=st.session_state.quote_data.get('numero_persone', 50))
        
        tipologia_servizio = st.text_input("Tipologia Servizio", 
                                         value=st.session_state.quote_data.get('tipologia_servizio', ''),
                                         placeholder="es. COCKTAIL di benvenuto e apericena")
        
        # Dettagli servizio
        st.subheader("Dettagli Servizio")
        
        col3, col4 = st.columns([1, 1])
        
        with col3:
            tipologia_buffet = st.selectbox("Tipologia Buffet", 
                                          ["Buffet e servizio", "Solo buffet", "Servizio al tavolo"],
                                          index=0)
            
            bicchieri = st.selectbox("Bicchieri", 
                                   ["Calici da vino in materiali ecocompatibili usa e getta", 
                                    "Bicchieri di vetro", 
                                    "Bicchieri di plastica riutilizzabili"])
            
            posate = st.selectbox("Posate", 
                                ["Acciaio leggero", "Acciaio inox", "Posate compostabili"])
        
        with col4:
            stoviglie = st.selectbox("Stoviglie", 
                                   ["Ceramica e materiali ecocompatibili usa e getta", 
                                    "Ceramica tradizionale", 
                                    "Piatti compostabili"])
            
            acqua = st.selectbox("Acqua", 
                               ["Esclusa / da definire", "Inclusa", "Su richiesta"])
            
            vini = st.selectbox("Vini", 
                              ["Esclusi / da definire", "Inclusi", "Su richiesta"])
        
        servizio = st.selectbox("Servizio", 
                              ["Incluso", "Escluso", "Parzialmente incluso"])
        
        # Salva i dati
        st.session_state.quote_data.update({
            'riferimento': riferimento,
            'destinatario': destinatario,
            'luogo': luogo,
            'data_evento': data_evento,
            'ora_evento': ora_evento,
            'numero_persone': numero_persone,
            'tipologia_servizio': tipologia_servizio,
            'tipologia_buffet': tipologia_buffet,
            'bicchieri': bicchieri,
            'posate': posate,
            'stoviglie': stoviglie,
            'acqua': acqua,
            'vini': vini,
            'servizio': servizio,
            'data_ora': f"{data_evento.strftime('%A %d %B %Y')} ore {ora_evento.strftime('%H,%M')}" if data_evento and ora_evento else ""
        })
    
    elif sezione == "🍽️ Menu":
        st.header("🍽️ Selezione Menu")
        
        # Carica elementi del menu dal file JSON
        all_menu_items = load_menu_items()
        categories = list(set([item['categoria'] for item in all_menu_items.values()]))
        
        # Filtra per categoria
        categoria_selezionata = st.selectbox("Filtra per categoria:", 
                                           ["Tutte"] + sorted(categories))
        
        # Mostra elementi disponibili
        st.subheader("Menu Disponibili")
        
        items_to_show = all_menu_items.items()
        if categoria_selezionata != "Tutte":
            items_to_show = [(k, v) for k, v in all_menu_items.items() 
                           if v['categoria'] == categoria_selezionata]
        
        for nome, dettagli in items_to_show:
            col1, col2 = st.columns([4, 1])
            
            with col1:
                st.write(f"**{nome}**")
                st.write(f"*{dettagli['categoria']}* - {dettagli['descrizione']}")
            
            with col2:
                if st.button("Aggiungi", key=f"add_{nome}"):
                    new_item = {
                        'nome': nome,
                        'categoria': dettagli['categoria'],
                        'descrizione': dettagli['descrizione']
                    }
                    # Controlla se non è già presente nel menu selezionato
                    if not any(item['nome'] == nome for item in st.session_state.menu_items):
                        st.session_state.menu_items.append(new_item)
                        st.success(f"Aggiunto: {nome}")
                    else:
                        st.warning(f"{nome} è già nel menu selezionato")
        
        # Aggiungi elemento personalizzato
        st.subheader("Aggiungi Nuovo Elemento al Database")
        
        with st.form("custom_item_form"):
            col1, col2 = st.columns([2, 1])
            
            with col1:
                custom_nome = st.text_input("Nome Piatto")
                custom_descrizione = st.text_area("Descrizione/Ingredienti")
            
            with col2:
                custom_categoria = st.selectbox("Categoria", 
                                              ["Aperitivo", "Antipasti", "Primi", "Secondi", 
                                               "Contorni", "Pasticceria Salata", "Pane", 
                                               "Dolci", "Frutta", "Bevande"])
            
            if st.form_submit_button("Aggiungi al Database"):
                if custom_nome:
                    # Aggiunge al database persistente
                    add_new_menu_item(custom_nome, custom_categoria, custom_descrizione)
                    st.success(f"Aggiunto '{custom_nome}' al database! Sarà disponibile per tutti i preventivi futuri.")
                    st.rerun()  # Ricarica la pagina per mostrare il nuovo elemento
        
        # Mostra menu corrente per questo preventivo
        if st.session_state.menu_items:
            st.subheader("Menu Selezionato per questo Preventivo")
            
            for i, item in enumerate(st.session_state.menu_items):
                col1, col2 = st.columns([4, 1])
                
                with col1:
                    st.write(f"**{item['nome']}**")
                    if item.get('descrizione'):
                        st.write(f"*{item['descrizione']}*")
                
                with col2:
                    if st.button("Rimuovi", key=f"remove_{i}"):
                        st.session_state.menu_items.pop(i)
                        st.rerun()
    
    elif sezione == "💰 Prezzi":
        st.header("💰 Calcolo Prezzi")
        
        col1, col2 = st.columns([1, 1])
        
        with col1:
            prezzo_persona = st.number_input("Prezzo per persona €", 
                                           min_value=0.0, 
                                           value=st.session_state.quote_data.get('prezzo_persona', 28.0),
                                           step=0.5)
            
            costo_cameriere = st.number_input("Costo Cameriere €", 
                                            min_value=0.0, 
                                            value=st.session_state.quote_data.get('costo_cameriere', 200.0),
                                            step=10.0)
        
        with col2:
            numero_persone = st.session_state.quote_data.get('numero_persone', 50)
            st.metric("Numero Persone", numero_persone)
            
            totale_base = prezzo_persona * numero_persone
            st.metric("Totale Base", f"€ {totale_base}")
            
            totale_finale = totale_base + costo_cameriere
            st.metric("Totale Finale", f"€ {totale_finale}")
        
        # Note aggiuntive
        note = st.text_area("Note Aggiuntive", 
                          value=st.session_state.quote_data.get('note', ''),
                          placeholder="Eventuali note aggiuntive per il cliente...")
        
        # Salva i dati
        st.session_state.quote_data.update({
            'prezzo_persona': prezzo_persona,
            'costo_cameriere': costo_cameriere,
            'note': note,
            'menu_items': st.session_state.menu_items
        })
        
        # Riassunto costi
        st.subheader("Riassunto Costi")
        
        # Create a simple summary without individual item prices
        st.info(f"**Menu selezionato:** {len(st.session_state.menu_items)} elementi")
        if st.session_state.menu_items:
            for item in st.session_state.menu_items:
                st.write(f"• {item['nome']}")
        
        df_costi = pd.DataFrame([
            ["Servizio per persona", f"€ {prezzo_persona}", f"{numero_persone} persone", f"€ {totale_base}"],
            ["Cameriere", f"€ {costo_cameriere}", "1", f"€ {costo_cameriere}"],
            ["**TOTALE**", "", "", f"**€ {totale_finale}**"]
        ], columns=["Descrizione", "Prezzo Unitario", "Quantità", "Totale"])
        
        st.table(df_costi)
        
        st.info("*IVA 22% esclusa")
    
    elif sezione == "📄 Anteprima":
        st.header("📄 Anteprima e Export")
        
        if not st.session_state.quote_data:
            st.warning("Completa prima i dati dell'evento nella sezione 'Dati Evento'")
            return
        
        # Anteprima
        st.subheader("Anteprima Preventivo")
        
        quote_data = st.session_state.quote_data
        
        st.markdown(f"**Ref:** {quote_data.get('riferimento', '')}")
        st.markdown(f"**Luogo:** {quote_data.get('luogo', '')}")
        st.markdown(f"**Data e Ora:** {quote_data.get('data_ora', '')}")
        st.markdown(f"**Tipologia servizio:** {quote_data.get('tipologia_servizio', '')}")
        
        st.markdown("---")
        
        st.markdown(f"**Gentile {quote_data.get('destinatario', '')},**")
        st.markdown(f"come da accordi Vi proponiamo la nostra migliore offerta per l'evento di cui in oggetto per numero {quote_data.get('numero_persone', '')} persone.")
        
        # Menu
        if st.session_state.menu_items:
            st.subheader("Menu Selezionato")
            
            # Group items by category for display
            menu_by_category = {}
            for item in st.session_state.menu_items:
                categoria = item.get('categoria', 'Altro')
                if categoria not in menu_by_category:
                    menu_by_category[categoria] = []
                menu_by_category[categoria].append(item)
            
            # Define category order
            category_order = [
                'Aperitivo', 'Antipasti', 'Primi', 'Secondi', 
                'Contorni', 'Pasticceria Salata', 'Pane', 
                'Dolci', 'Frutta', 'Bevande', 'Altro'
            ]
            
            # Display grouped menu
            for categoria in category_order:
                if categoria in menu_by_category and menu_by_category[categoria]:
                    st.markdown(f"### **{categoria}**")
                    for item in menu_by_category[categoria]:
                        st.markdown(f"**{item['nome'].upper()}**")
                        if item.get('descrizione'):
                            st.markdown(f"*{item['descrizione']}*")
                        st.markdown("")
        
        # Totali
        numero_persone = quote_data.get('numero_persone', 0)
        prezzo_persona = quote_data.get('prezzo_persona', 0)
        costo_cameriere = quote_data.get('costo_cameriere', 0)
        totale_base = prezzo_persona * numero_persone
        totale_finale = totale_base + costo_cameriere
        
        st.markdown("---")
        st.markdown("**Valorizzazione dell'offerta:**")
        st.markdown(f"Il costo del servizio è di {prezzo_persona} Euro/persona per un minimo di {numero_persone} persone.")
        st.markdown(f"**Totale {numero_persone} convenuti: € {totale_base}**")
        if costo_cameriere > 0:
            st.markdown(f"**Cameriere: € {costo_cameriere}**")
        st.markdown(f"## **TOTALE: € {totale_finale}**")
        st.markdown("*(*) Escluso IVA 22%*")
        
        if quote_data.get('note'):
            st.markdown("---")
            st.markdown("**Note Aggiuntive:**")
            st.markdown(quote_data['note'])
        
        # Export
        st.subheader("Esporta Preventivo")
        
        col1, col2 = st.columns([1, 1])
        
        with col1:
            if st.button("📄 Scarica Word (.docx)", type="primary"):
                doc = create_word_document(quote_data)
                
                # Salva in memoria
                file_stream = io.BytesIO()
                doc.save(file_stream)
                file_stream.seek(0)
                
                filename = f"Preventivo_{quote_data.get('riferimento', 'Cliente')}_{date.today().strftime('%Y%m%d')}.docx"
                
                st.download_button(
                    label="⬇️ Download Documento",
                    data=file_stream.getvalue(),
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        
        with col2:
            if st.button("🔄 Nuovo Preventivo"):
                # Reset dati
                st.session_state.quote_data = {}
                st.session_state.menu_items = []
                st.success("Dati resettati. Puoi iniziare un nuovo preventivo.")
                st.rerun()

if __name__ == "__main__":
    main()
